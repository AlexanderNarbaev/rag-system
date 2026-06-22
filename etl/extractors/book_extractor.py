"""Book extractor supporting EPUB, PDF, DOCX formats."""

import logging
import re
from collections.abc import Generator
from pathlib import Path
from typing import Any

from etl.extractors.base_extractor import BaseExtractor, ExtractedDocument, ExtractorConfig

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)


class BookExtractor(BaseExtractor):
    """Extracts content from EPUB, PDF, and DOCX books with chapter/section detection."""

    EPUB_MIME = "application/epub+zip"
    PDF_MIME = "application/pdf"
    DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def __init__(self, config: ExtractorConfig):
        super().__init__(config)
        self._validated = False
        self._source_files: list[Path] = []

    async def validate_connection(self) -> bool:
        """Validate that source files exist and are accessible."""
        base_path = Path(self.config.base_url) if self.config.base_url else None
        if not base_path:
            logger.error("No base_url provided for book extractor")
            return False
        if not base_path.exists():
            logger.error(f"Book source directory does not exist: {base_path}")
            return False

        supported_extensions = {".epub", ".pdf", ".docx"}
        self._source_files = []
        for ext in supported_extensions:
            self._source_files.extend(list(base_path.rglob(f"*{ext}")))
            self._source_files.extend(list(base_path.rglob(f"*{ext.upper()}")))

        if self.config.exclude_patterns:
            self._source_files = [
                f for f in self._source_files if not any(p in str(f) for p in self.config.exclude_patterns)
            ]

        self._validated = len(self._source_files) > 0
        if self._validated:
            logger.info(f"Found {len(self._source_files)} book files in {base_path}")
        else:
            logger.warning(f"No supported book files found in {base_path}")
        return self._validated

    async def extract(self) -> Generator[ExtractedDocument, None, None]:
        """Extract documents from all supported book files."""
        if not self._validated:
            await self.validate_connection()

        for file_path in self._source_files:
            try:
                ext = file_path.suffix.lower()
                if ext == ".epub":
                    docs = await self._extract_epub(file_path)
                elif ext == ".pdf":
                    docs = await self._extract_pdf(file_path)
                elif ext == ".docx":
                    docs = await self._extract_docx(file_path)
                else:
                    continue

                for doc in docs:
                    yield doc
            except Exception as e:
                logger.error(f"Failed to extract {file_path}: {e}", exc_info=True)

    async def _extract_epub(self, file_path: Path) -> list[ExtractedDocument]:
        """Extract content from an EPUB file."""
        try:
            from ebooklib import ITEM_DOCUMENT, epub
        except ImportError:
            logger.warning("ebooklib not installed, skipping EPUB extraction")
            return []

        book = epub.read_epub(str(file_path))
        title = self._get_epub_metadata(book, "title") or file_path.stem
        author = self._get_epub_metadata(book, "creator")
        isbn = self._get_epub_identifier(book)

        documents = []
        chapters: list[dict[str, Any]] = []
        section_counter = 0

        for item in book.get_items_of_type(ITEM_DOCUMENT):
            try:
                content = item.get_content().decode("utf-8")
            except UnicodeDecodeError:
                content = item.get_content().decode("latin-1", errors="replace")

            cleaned = self._clean_html(content)
            heading_info = self._extract_headings(cleaned)

            if heading_info:
                for heading_text, heading_level, section_text in heading_info:
                    section_counter += 1
                    section_title = heading_text or f"Section {section_counter}"
                    chapters.append(
                        {
                            "title": section_title,
                            "content": section_text,
                            "level": heading_level,
                            "index": section_counter,
                        }
                    )
            else:
                section_counter += 1
                chapters.append(
                    {
                        "title": f"Chapter {section_counter}",
                        "content": cleaned,
                        "level": 1,
                        "index": section_counter,
                    }
                )

        for ch in chapters:
            doc = ExtractedDocument(
                source_id=f"book_epub_{file_path.stem}_ch{ch['index']}",
                source_type="book",
                title=f"{title} — {ch['title']}",
                content=ch["content"],
                content_type="text",
                metadata={
                    "book_title": title,
                    "author": author,
                    "isbn": isbn,
                    "format": "epub",
                    "chapter": ch["title"],
                    "chapter_level": ch["level"],
                    "chapter_index": ch["index"],
                    "file_path": str(file_path),
                },
                access_level=self.config.source_type if self.config.source_type else "internal",
            )
            documents.append(doc)

        return documents

    async def _extract_pdf(self, file_path: Path) -> list[ExtractedDocument]:
        """Extract text content from a PDF file."""
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.warning("pypdf not installed, skipping PDF extraction")
            return []

        reader = PdfReader(str(file_path))
        metadata = reader.metadata or {}
        title = metadata.get("/Title") or file_path.stem
        author = metadata.get("/Author", "")

        full_text_parts = []
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            full_text_parts.append(text)

        full_text = "\n\n".join(full_text_parts)

        sections = self._split_into_sections(full_text)
        documents = []
        for idx, (section_title, section_text) in enumerate(sections):
            doc = ExtractedDocument(
                source_id=f"book_pdf_{file_path.stem}_s{idx}",
                source_type="book",
                title=f"{title} — {section_title or f'Section {idx + 1}'}",
                content=section_text,
                content_type="text",
                metadata={
                    "book_title": title,
                    "author": author,
                    "format": "pdf",
                    "total_pages": len(reader.pages),
                    "section_title": section_title,
                    "section_index": idx,
                    "file_path": str(file_path),
                },
            )
            documents.append(doc)

        return documents

    async def _extract_docx(self, file_path: Path) -> list[ExtractedDocument]:
        """Extract content from a DOCX file."""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed, skipping DOCX extraction")
            return []

        docx = Document(str(file_path))
        title = ""
        author = ""

        if docx.core_properties:
            title = docx.core_properties.title or file_path.stem
            author = docx.core_properties.author or ""

        paragraphs_text = []
        current_chapter_title = ""
        chapter_text = ""
        chapters: list[tuple[str, str]] = []

        for para in docx.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            if style_name and any(
                level in style_name.lower() for level in ("heading", "title", "heading1", "heading2", "heading3")
            ):
                if chapter_text.strip():
                    chapters.append((current_chapter_title or title, chapter_text.strip()))
                current_chapter_title = text
                chapter_text = ""
            else:
                chapter_text += text + "\n"
                paragraphs_text.append(text)

        if chapter_text.strip():
            chapters.append((current_chapter_title or title, chapter_text.strip()))

        if not chapters:
            chapters.append((title, "\n\n".join(paragraphs_text)))

        documents = []
        for idx, (ch_title, ch_content) in enumerate(chapters):
            doc = ExtractedDocument(
                source_id=f"book_docx_{file_path.stem}_ch{idx}",
                source_type="book",
                title=f"{title} — {ch_title}",
                content=ch_content,
                content_type="text",
                metadata={
                    "book_title": title,
                    "author": author,
                    "format": "docx",
                    "chapter": ch_title,
                    "chapter_index": idx,
                    "file_path": str(file_path),
                },
            )
            documents.append(doc)

        return documents

    def should_process(self, doc: ExtractedDocument, last_hash: str) -> bool:
        """Check if document needs processing based on content hash."""
        if not last_hash:
            return True
        return self.compute_hash(doc.content) != last_hash

    @staticmethod
    def _get_epub_metadata(book, field: str) -> str:
        try:
            for item in book.get_metadata("DC", field):
                return item[0]
        except Exception:
            pass
        return ""

    @staticmethod
    def _get_epub_identifier(book) -> str:
        try:
            for item in book.get_metadata("DC", "identifier"):
                val = item[0]
                if "isbn" in str(val).lower() or re.match(r"[\d-]{10,17}", str(val)):
                    return str(val)
        except Exception:
            pass
        return ""

    @staticmethod
    def _clean_html(html_content: str) -> str:
        """Remove HTML tags, scripts, and styles."""
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(html_content, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            text = soup.get_text(separator="\n")
        except ImportError:
            text = re.sub(r"<script[^>]*>.*?</script>", "", html_content, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r"<[^>]+>", "", text)
            text = re.sub(r"&nbsp;", " ", text)
            text = re.sub(r"&amp;", "&", text)
            text = re.sub(r"&lt;", "<", text)
            text = re.sub(r"&gt;", ">", text)

        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)

    @staticmethod
    def _extract_headings(text: str) -> list[tuple[str, int, str]]:
        """Extract headings and their section content from HTML text."""
        heading_pattern = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)
        matches = list(heading_pattern.finditer(text))

        if not matches:
            return []

        sections = []
        for i, match in enumerate(matches):
            level = len(match.group(1))
            heading_text = match.group(2).strip()
            start = match.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            section_text = text[start:end].strip()
            sections.append((heading_text, level, section_text))

        return sections

    @staticmethod
    def _split_into_sections(text: str) -> list[tuple[str, str]]:
        """Split text into sections by common chapter patterns."""
        chapter_patterns = [
            r"^(?:Chapter|CHAPTER|Глава|ГЛАВА)\s+\d+",
            r"^(?:Section|SECTION|Раздел)\s+\d+",
            r"^\d+\.\s+\w",
            r"^(?:Part|PART|Часть)\s+\d+",
            r"^[IVXLC]+\.\s+\w",
        ]

        combined = "|".join(f"(?:{p})" for p in chapter_patterns)
        section_re = re.compile(combined, re.MULTILINE)

        splits = list(section_re.finditer(text))
        if not splits:
            return [("", text)]

        sections = []
        for i, match in enumerate(splits):
            title = match.group(0).strip()
            start = match.start()
            end = splits[i + 1].start() if i + 1 < len(splits) else len(text)
            section_text = text[start:end].strip()
            sections.append((title, section_text))

        return sections
